import docx #python-docx

d=docx.Document('demo.docx')
print(d.paragraphs)



print(d.paragraphs[0].text)


print("\n")


p=d.paragraphs[1]

print(p.runs[1].text)
print(p.runs[1].bold)
print(p.runs[1].italic)


p.runs[3].undeline=True
p.runs[3].text='Hey there['



print(p.text)
print(p.style)
p.style='Title'

d.save('demo2.docx')


d=docx.Document()
d.add_paragraph('Gello This is a para');
d.add_paragraph('Gello This is a para1');
d.add_paragraph('Gello This is a para22');

p=d.paragraphs[0]
p.add_run('This is a new run')
print(p.runs)
p.runs[1].bold=True

d.save('new.docx')
